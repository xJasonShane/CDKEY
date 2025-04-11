import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import random
import string

class GenerateTab:
    def __init__(self, parent):
        self.parent = parent
        self.setup_ui()
    
    def setup_ui(self):
        """设置生成页面UI"""
        # 左侧控制面板
        self.control_frame = ttk.LabelFrame(self.parent, text="生成设置")
        self.control_frame.pack(side=tk.LEFT, fill=tk.Y, padx=5, pady=5)
        
        # 位数设置
        ttk.Label(self.control_frame, text="CDK位数:").grid(row=0, column=0, sticky=tk.W)
        self.length_entry = ttk.Entry(self.control_frame)
        self.length_entry.grid(row=0, column=1, padx=5, pady=5)
        
        # 数量设置
        ttk.Label(self.control_frame, text="生成数量:").grid(row=1, column=0, sticky=tk.W)
        self.count_entry = ttk.Entry(self.control_frame)
        self.count_entry.grid(row=1, column=1, padx=5, pady=5)
        
        # 字符组成选项
        self.use_digits = tk.BooleanVar(value=True)
        ttk.Checkbutton(self.control_frame, text="包含数字", variable=self.use_digits).grid(row=2, column=0, columnspan=2, sticky=tk.W)
        
        self.use_letters = tk.BooleanVar(value=True)
        ttk.Checkbutton(self.control_frame, text="包含字母", variable=self.use_letters).grid(row=3, column=0, columnspan=2, sticky=tk.W)
        
        self.use_special = tk.BooleanVar(value=False)
        ttk.Checkbutton(self.control_frame, text="包含特殊字符", variable=self.use_special).grid(row=4, column=0, columnspan=2, sticky=tk.W)
        
        self.case_sensitive = tk.BooleanVar(value=False)
        ttk.Checkbutton(self.control_frame, text="区分大小写", variable=self.case_sensitive).grid(row=5, column=0, columnspan=2, sticky=tk.W)
        
        # 生成按钮
        ttk.Button(self.control_frame, text="生成", command=self.generate_cdkeys).grid(row=6, column=0, columnspan=2, pady=10)
        
        # 右侧预览面板
        self.preview_frame = ttk.LabelFrame(self.parent, text="预览")
        self.preview_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        self.preview_text = tk.Text(self.preview_frame, wrap=tk.WORD)
        self.preview_text.pack(fill=tk.BOTH, expand=True)
        
        # 导出按钮
        ttk.Button(self.preview_frame, text="导出", command=self.export_cdkeys).pack(pady=5)
    
    def generate_cdkeys(self):
        """生成CDKEY"""
        try:
            length = int(self.length_entry.get())
            count = int(self.count_entry.get())
            
            if length <= 0 or count <= 0:
                messagebox.showerror("错误", "位数和数量必须大于0")
                return
            
            # 清空预览
            self.preview_text.delete(1.0, tk.END)
            
            # 生成字符池
            chars = ""
            if self.use_digits.get():
                chars += string.digits
            if self.use_letters.get():
                if self.case_sensitive.get():
                    chars += string.ascii_letters
                else:
                    chars += string.ascii_uppercase
            if self.use_special.get():
                chars += "!@#$%^&*()"
            
            if not chars:
                messagebox.showerror("错误", "至少选择一种字符类型")
                return
            
            # 生成CDKEY
            cdkeys = []
            for _ in range(count):
                cdkey = ''.join(random.choice(chars) for _ in range(length))
                cdkeys.append(cdkey)
                self.preview_text.insert(tk.END, cdkey + "\n")
            
            self.generated_cdkeys = cdkeys
            messagebox.showinfo("成功", f"已生成{count}个{length}位CDKEY")
        except ValueError:
            messagebox.showerror("错误", "请输入有效的数字")
    
    def export_cdkeys(self):
        """导出CDKEY"""
        if not hasattr(self, 'generated_cdkeys') or not self.generated_cdkeys:
            messagebox.showerror("错误", "请先生成CDKEY")
            return
        
        # 弹出文件保存对话框
        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[
                ("文本文件", "*.txt"),
                ("Excel文件", "*.xlsx"),
                ("Word文件", "*.docx"),
                ("所有文件", "*.*")
            ],
            title="保存CDKEY文件"
        )
        
        if not file_path:  # 用户取消选择
            return
            
        try:
            if file_path.endswith('.xlsx'):
                from openpyxl import Workbook
                wb = Workbook()
                ws = wb.active
                ws.title = "CDKEY列表"
                for i, cdkey in enumerate(self.generated_cdkeys, 1):
                    ws.cell(row=i, column=1, value=cdkey)
                wb.save(file_path)
            elif file_path.endswith('.docx'):
                from docx import Document
                doc = Document()
                doc.add_heading('CDKEY列表', 0)
                for cdkey in self.generated_cdkeys:
                    doc.add_paragraph(cdkey)
                doc.save(file_path)
            else:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write('\n'.join(self.generated_cdkeys))
            messagebox.showinfo("成功", f"CDKEY已成功导出到 {file_path}")
        except Exception as e:
            messagebox.showerror("错误", f"导出失败: {str(e)}")